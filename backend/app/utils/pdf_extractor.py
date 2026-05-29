import os
from typing import Optional
from app.utils.logger import app_logger


def extract_text_from_file(file_path: str) -> str:
    """Extract text from PDF, DOCX, or TXT files."""
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return _extract_from_pdf(file_path)
    elif ext in (".docx", ".doc"):
        return _extract_from_docx(file_path)
    elif ext == ".txt":
        return _extract_from_txt(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def _extract_from_pdf(file_path: str) -> str:
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
        return "\n".join(text_parts)
    except Exception as e:
        app_logger.warning(f"pdfplumber failed, trying PyPDF2: {e}")
        return _extract_pdf_fallback(file_path)


def _extract_pdf_fallback(file_path: str) -> str:
    try:
        import PyPDF2
        text_parts = []
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
        return "\n".join(text_parts)
    except Exception as e:
        app_logger.error(f"PDF extraction failed: {e}")
        raise RuntimeError(f"Could not extract text from PDF: {e}")


def _extract_from_docx(file_path: str) -> str:
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        return "\n".join(paragraphs)
    except Exception as e:
        app_logger.error(f"DOCX extraction failed: {e}")
        raise RuntimeError(f"Could not extract text from DOCX: {e}")


def _extract_from_txt(file_path: str) -> str:
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception as e:
        app_logger.error(f"TXT extraction failed: {e}")
        raise RuntimeError(f"Could not read text file: {e}")
